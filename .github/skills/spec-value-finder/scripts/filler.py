"""記入シート例（テンプレート）を元に、確定値を埋めた新規ファイルを生成する。

Claude が候補を吟味して作成した findings.json を入力に取り、テンプレートを
複製して値・出典・確信度・要確認フラグを書き込む。

- Excel テンプレート: 「項目」列の項目名と findings.target を突合して各行へ記入。
                      未マッチの findings は末尾へ追記する。
- Word  テンプレート: 本文/表中の `{{項目名}}` プレースホルダを値で置換する。
"""
from __future__ import annotations

import json
from pathlib import Path

from models import normalize

CONFIDENCE_DISPLAY = {
    "high": "高", "medium": "中", "low": "低",
    "高": "高", "中": "中", "低": "低", "": "",
}

DEFAULT_COLUMNS = {
    "item": "項目", "value": "値", "source": "出典",
    "confidence": "確信度", "review": "要確認",
}


def load_findings(path: Path | str) -> list[dict]:
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    items = data.get("items") if isinstance(data, dict) else data
    if not isinstance(items, list):
        raise ValueError("findings は items リストである必要があります。")
    return items


def _conf(value: str) -> str:
    return CONFIDENCE_DISPLAY.get(str(value or "").strip().lower()
                                  if str(value).strip().lower() in CONFIDENCE_DISPLAY
                                  else str(value or "").strip(), str(value or ""))


def _review_text(needs_review) -> str:
    return "要確認" if needs_review else ""


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------

def fill_excel(template: Path, findings: list[dict], out: Path,
               columns: dict[str, str] | None = None,
               sheet: str | None = None) -> dict:
    import openpyxl

    cols = {**DEFAULT_COLUMNS, **(columns or {})}
    wb = openpyxl.load_workbook(template)
    ws = wb[sheet] if sheet else wb.active

    # ヘッダ行を探す（item 見出しを含む最初の行）
    header_row = None
    col_idx: dict[str, int] = {}
    for row in ws.iter_rows():
        labels = {str(c.value).strip(): c.column for c in row if c.value is not None}
        if cols["item"] in labels:
            header_row = row[0].row
            for key, name in cols.items():
                if name in labels:
                    col_idx[key] = labels[name]
            break
    if header_row is None:
        raise ValueError(
            f"テンプレートに項目列見出し '{cols['item']}' が見つかりません。"
            f" --col-item で見出し名を指定してください。")
    if "value" not in col_idx:
        raise ValueError(f"テンプレートに値列見出し '{cols['value']}' が見つかりません。")

    by_target = {normalize(f.get("target", "")): f for f in findings}
    written: set[str] = set()

    def write_row(r: int, f: dict) -> None:
        ws.cell(row=r, column=col_idx["value"]).value = f.get("value", "")
        if "source" in col_idx:
            ws.cell(row=r, column=col_idx["source"]).value = f.get("source", "")
        if "confidence" in col_idx:
            ws.cell(row=r, column=col_idx["confidence"]).value = _conf(f.get("confidence", ""))
        if "review" in col_idx:
            ws.cell(row=r, column=col_idx["review"]).value = _review_text(f.get("needs_review"))

    for row in ws.iter_rows(min_row=header_row + 1):
        item_cell = row[col_idx["item"] - 1] if col_idx["item"] - 1 < len(row) else None
        if item_cell is None or item_cell.value is None:
            continue
        key = normalize(str(item_cell.value))
        if key in by_target:
            write_row(item_cell.row, by_target[key])
            written.add(key)

    # 未マッチの findings を末尾へ追記
    appended = 0
    next_row = ws.max_row + 1
    for f in findings:
        key = normalize(f.get("target", ""))
        if key in written or not key:
            continue
        ws.cell(row=next_row, column=col_idx["item"]).value = f.get("target", "")
        write_row(next_row, f)
        next_row += 1
        appended += 1

    Path(out).parent.mkdir(parents=True, exist_ok=True)
    wb.save(out)
    return {
        "out": str(out),
        "matched": len(written),
        "appended": appended,
        "unmatched_findings": [f.get("target") for f in findings
                               if normalize(f.get("target", "")) not in written],
    }


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> int:
    text = paragraph.text
    if "{{" not in text:
        return 0
    new_text = text
    count = 0
    for target, value in mapping.items():
        token = "{{" + target + "}}"
        if token in new_text:
            count += new_text.count(token)
            new_text = new_text.replace(token, value)
    if count and new_text != text:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
    return count


def fill_word(template: Path, findings: list[dict], out: Path) -> dict:
    import docx

    doc = docx.Document(str(template))
    mapping = {str(f.get("target", "")): str(f.get("value", "")) for f in findings}
    replaced = 0
    for para in doc.paragraphs:
        replaced += _replace_in_paragraph(para, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replaced += _replace_in_paragraph(para, mapping)

    Path(out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return {"out": str(out), "replaced": replaced}


def fill(template: Path | str, findings: list[dict], out: Path | str,
         columns: dict[str, str] | None = None, sheet: str | None = None) -> dict:
    template = Path(template)
    suffix = template.suffix.lower()
    if suffix in (".xlsx", ".xlsm"):
        return fill_excel(template, findings, Path(out), columns=columns, sheet=sheet)
    if suffix == ".docx":
        return fill_word(template, findings, Path(out))
    raise ValueError(f"対応していないテンプレート形式です: {suffix}")
